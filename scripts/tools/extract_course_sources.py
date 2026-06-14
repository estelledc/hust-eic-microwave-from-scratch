"""Extract a curated source index and publishable teaching images.

This script treats ``sources/`` as local-only input. It never copies source
PDF/DOCX files into the published site. Outputs are intentionally small:

- docs/audit/SOURCE_EXTRACTION_AUDIT.md
- docs/audit/source_extraction_index.json
- assets/images/course/*-source-preview.webp
- assets/images/exp2/*.webp
"""
from __future__ import annotations

import json
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import fitz
from docx import Document
from PIL import Image, ImageOps


ROOT = Path(__file__).resolve().parents[2]
PDF_SOURCE_DIR = ROOT / "sources" / "课件PDF"
EXP2_SOURCE_DIR = ROOT / "sources" / "微波实验2"
COURSE_IMAGE_DIR = ROOT / "assets" / "images" / "course"
EXP2_IMAGE_DIR = ROOT / "assets" / "images" / "exp2"
AUDIT_MD = ROOT / "docs" / "audit" / "SOURCE_EXTRACTION_AUDIT.md"
AUDIT_JSON = ROOT / "docs" / "audit" / "source_extraction_index.json"
INTEGRATION_REVIEW_MD = ROOT / "docs" / "audit" / "COURSE_INTEGRATION_REVIEW.md"


@dataclass(frozen=True)
class CourseSource:
    filename: str
    slug: str
    topic: str
    stages: str
    targets: tuple[str, ...]
    keywords: tuple[str, ...]


@dataclass(frozen=True)
class PageRange:
    start: int
    end: int
    summary: str
    treatment: str
    targets: tuple[str, ...]


COURSE_SOURCES: tuple[CourseSource, ...] = (
    CourseSource(
        "传输线理论1.pdf",
        "transmission-line-theory",
        "传输线理论、长线模型、反射与驻波",
        "knowledge/01",
        (
            "content/knowledge/01-传播与传输线/README.md",
            "content/knowledge/01-传播与传输线/99-自检清单与常见误区.md",
        ),
        ("传输线", "特性阻抗", "反射系数", "驻波", "输入阻抗"),
    ),
    CourseSource(
        "传输线理论 - 圆图.pdf",
        "smith-chart",
        "Smith 圆图、归一化阻抗、沿线旋转",
        "knowledge/02",
        (
            "content/knowledge/02-反射与匹配/02-Smith圆图怎么读.md",
            "content/knowledge/02-反射与匹配/99-自检清单与常见误区.md",
        ),
        ("Smith", "圆图", "归一化", "导纳", "驻波比"),
    ),
    CourseSource(
        "阻抗匹配实例.pdf",
        "impedance-matching-examples",
        "阻抗匹配实例、支节匹配与工程流程",
        "knowledge/02",
        (
            "content/knowledge/02-反射与匹配/03-并联支节匹配.md",
            "content/knowledge/02-反射与匹配/99-自检清单与常见误区.md",
        ),
        ("阻抗匹配", "支节", "Smith", "匹配", "导纳"),
    ),
    CourseSource(
        "矩形波导1.pdf",
        "rectangular-waveguide",
        "矩形波导、截止、模式、单模窗口",
        "knowledge/03-05",
        (
            "content/knowledge/03-波导中的场与边界/README.md",
            "content/knowledge/04-截止色散与速度/README.md",
            "content/knowledge/05-矩形波导工程计算/README.md",
        ),
        ("矩形波导", "截止", "TE", "TM", "单模"),
    ),
    CourseSource(
        "圆波导.pdf",
        "circular-waveguide",
        "圆波导、贝塞尔根、主模与单模边界",
        "knowledge/06",
        (
            "content/knowledge/06-圆波导同轴线微带线/01-圆波导模式与贝塞尔根.md",
            "content/knowledge/06-圆波导同轴线微带线/99-自检清单与常见误区.md",
        ),
        ("圆波导", "贝塞尔", "TE11", "TM01", "截止"),
    ),
    CourseSource(
        "同轴线.pdf",
        "coaxial-line",
        "同轴线 TEM、高阶模上限、结构选择",
        "knowledge/06",
        (
            "content/knowledge/06-圆波导同轴线微带线/02-同轴线TEM与高阶模.md",
            "content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md",
        ),
        ("同轴线", "TEM", "高阶模", "截止", "特性阻抗"),
    ),
    CourseSource(
        "微波传输线1-2.pdf",
        "microwave-transmission-lines",
        "微波传输线类型与导波结构对照",
        "knowledge/06",
        (
            "content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md",
            "content/knowledge/README.md",
        ),
        ("微波传输线", "矩形波导", "圆波导", "同轴线", "微带线"),
    ),
    CourseSource(
        "微波集成传输线.pdf",
        "integrated-microwave-lines",
        "微带线、集成传输线、准 TEM 与有效介电常数",
        "knowledge/06",
        (
            "content/knowledge/06-圆波导同轴线微带线/03-微带线准TEM与有效介电常数.md",
            "content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md",
        ),
        ("微带线", "集成", "准TEM", "有效介电常数", "带状线"),
    ),
    CourseSource(
        "微波谐振腔.pdf",
        "microwave-resonator",
        "微波谐振器、谐振腔、品质因数",
        "knowledge/07-08",
        (
            "content/knowledge/07-实验测量与微波元件/03-谐振器Q值与功率传输法.md",
            "content/knowledge/08-谐振器网络与课程综合/01-微波谐振器与谐振腔.md",
        ),
        ("谐振", "谐振腔", "品质因数", "Q", "模式"),
    ),
    CourseSource(
        "微波网络基础2.pdf",
        "microwave-network",
        "微波网络参数、S 参数、测量读数",
        "knowledge/07-08",
        (
            "content/knowledge/07-实验测量与微波元件/01-S参数与矢量网络分析仪.md",
            "content/knowledge/08-谐振器网络与课程综合/02-微波网络基础与S参数.md",
        ),
        ("微波网络", "S参数", "散射参数", "二端口", "矩阵"),
    ),
)


COURSE_PAGE_RANGES: dict[str, tuple[PageRange, ...]] = {
    "transmission-line-theory": (
        PageRange(1, 2, "封面、教学目的、核心知识点。只用于确认课程主线，不进入正文。", "审计", ()),
        PageRange(3, 16, "微波波段、长线/短线、分布参数、场方法与路方法的区别。", "融入", ("content/knowledge/README.md", "content/knowledge/01-传播与传输线/01-长线短线与分布参数.md")),
        PageRange(17, 23, "传输线作为导波系统的分类，TEM、TE/TM、表面波，以及横向问题/纵向问题的分工。", "融入", ("content/knowledge/01-传播与传输线/README.md", "content/knowledge/03-波导中的场与边界/00-从传输线到波导.md", "content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md")),
        PageRange(24, 35, "分布参数小线元、电报方程、无耗波动方程、正反向行波、相移常数、相速和特性阻抗。", "融入", ("content/knowledge/01-传播与传输线/01-长线短线与分布参数.md", "content/knowledge/01-传播与传输线/02-行波相位常数与特性阻抗.md")),
        PageRange(36, 46, "接负载后的电压电流表达式、输入阻抗和输入导纳。", "融入", ("content/knowledge/01-传播与传输线/03-反射驻波与输入阻抗.md", "content/knowledge/02-反射与匹配/01-多段线并联与四分之一波长.md")),
        PageRange(47, 57, "反射系数、波腹波节、驻波比及驻波比约束例题。", "融入", ("content/knowledge/01-传播与传输线/03-反射驻波与输入阻抗.md", "content/knowledge/02-反射与匹配/02-Smith圆图怎么读.md")),
        PageRange(58, 62, "行波状态：匹配负载、无反射、沿线阻抗恒为特性阻抗。", "融入", ("content/knowledge/01-传播与传输线/04-行波纯驻波与行驻波.md",)),
        PageRange(63, 76, "纯驻波状态：短路、开路、纯电抗负载及 λ/4 阻抗性质变换。", "融入", ("content/knowledge/01-传播与传输线/05-开短路线周期性与测量.md", "content/knowledge/02-反射与匹配/01-多段线并联与四分之一波长.md")),
        PageRange(77, 85, "行驻波状态、任意复负载、波腹/波节位置和 Rmax/Rmin。", "融入", ("content/knowledge/01-传播与传输线/04-行波纯驻波与行驻波.md", "content/knowledge/01-传播与传输线/05-开短路线周期性与测量.md")),
        PageRange(86, 91, "λ/4 输入阻抗例题、特性阻抗优化使驻波比最小的例题。", "融入", ("content/knowledge/02-反射与匹配/01-多段线并联与四分之一波长.md", "content/knowledge/02-反射与匹配/99-自检清单与常见误区.md")),
        PageRange(92, 92, "结束页。", "审计", ()),
    ),
    "smith-chart": (
        PageRange(1, 3, "封面、教学目标和长线问题回顾。", "审计", ()),
        PageRange(4, 12, "Smith 圆图用途、等反射系数圆、沿线移动方向和电长度读数。", "融入", ("content/knowledge/02-反射与匹配/02-Smith圆图怎么读.md",)),
        PageRange(13, 17, "归一化阻抗、等电阻圆、等电抗圆及阻抗圆图构成。", "融入", ("content/knowledge/02-反射与匹配/02-Smith圆图怎么读.md",)),
        PageRange(18, 21, "开路/短路/匹配点、三条特殊线、上/下半圆含义、四个读数量。", "融入", ("content/knowledge/02-反射与匹配/02-Smith圆图怎么读.md", "content/knowledge/02-反射与匹配/99-自检清单与常见误区.md")),
        PageRange(22, 27, "已知负载求 SWR、波腹/波节和输入阻抗；由 SWR 与节点位置反推负载。", "融入", ("content/knowledge/02-反射与匹配/02-Smith圆图怎么读.md", "content/knowledge/01-传播与传输线/05-开短路线周期性与测量.md")),
        PageRange(28, 33, "导纳圆图、阻抗/导纳互换、λ/4 对称关系和并联题读图口径。", "融入", ("content/knowledge/02-反射与匹配/03-并联支节匹配.md", "content/knowledge/02-反射与匹配/99-自检清单与常见误区.md")),
        PageRange(34, 38, "双导线归一化输入导纳例题，多种圆图解法对比。", "融入", ("content/knowledge/02-反射与匹配/03-并联支节匹配.md",)),
        PageRange(39, 39, "结束页。", "审计", ()),
    ),
    "impedance-matching-examples": (
        PageRange(1, 1, "封面页。", "审计", ()),
        PageRange(2, 4, "微带贴片天线失配问题、等效负载、阻抗变换与导纳求和。", "融入", ("content/knowledge/02-反射与匹配/README.md", "content/knowledge/02-反射与匹配/03-并联支节匹配.md")),
        PageRange(5, 11, "λ/4 阻抗变换器原理、枝节补偿虚部、Rogers 5880 工程结果。", "融入", ("content/knowledge/02-反射与匹配/01-多段线并联与四分之一波长.md", "content/knowledge/02-反射与匹配/03-并联支节匹配.md")),
        PageRange(12, 14, "失配问题与等效负载的重复引入。", "审计", ("content/knowledge/02-反射与匹配/03-并联支节匹配.md",)),
        PageRange(15, 21, "单枝节匹配原理、接入距离 d、枝节长度 l 和反射降低结果。", "融入", ("content/knowledge/02-反射与匹配/03-并联支节匹配.md",)),
        PageRange(22, 24, "失配问题与等效负载的第三次引入。", "审计", ("content/knowledge/02-反射与匹配/03-并联支节匹配.md",)),
        PageRange(25, 31, "双枝节匹配原理、死区提醒、l1/l2 求解和反射降低结果。", "融入", ("content/knowledge/02-反射与匹配/03-并联支节匹配.md", "content/knowledge/02-反射与匹配/99-自检清单与常见误区.md")),
    ),
    "rectangular-waveguide": (
        PageRange(1, 2, "封面、教学目的。", "审计", ()),
        PageRange(3, 3, "矩形波导坐标、宽边 a 与窄边 b。", "融入", ("content/knowledge/03-波导中的场与边界/00-从传输线到波导.md", "content/knowledge/05-矩形波导工程计算/00-工程计算路线图.md")),
        PageRange(4, 6, "TEmn/TMmn 截止波数、截止波长、导波波长、主模与简并模。", "融入", ("content/knowledge/03-波导中的场与边界/03-金属边界与截止.md", "content/knowledge/04-截止色散与速度/01-三种波长.md", "content/knowledge/05-矩形波导工程计算/01-模谱主模与简并.md")),
        PageRange(7, 12, "模式截止图、可传输模例题、相速群速、波阻抗和工作波型图。", "融入", ("content/knowledge/04-截止色散与速度/02-色散相速与群速.md", "content/knowledge/05-矩形波导工程计算/04-可传输模判定与枚举.md")),
        PageRange(13, 18, "TE10 主模场分量、场分布和 TE10 传输特性。", "融入", ("content/knowledge/03-波导中的场与边界/04-从纵向场到全场.md", "content/knowledge/05-矩形波导工程计算/03-导波波长相速群速算例.md")),
        PageRange(19, 24, "TE10 功率容量、损耗、衰减和矩形波导尺寸选择原则。", "融入", ("content/knowledge/05-矩形波导工程计算/02-单模工作区与介质填充.md", "content/knowledge/05-矩形波导工程计算/99-自检清单与常见误区.md")),
        PageRange(25, 30, "TE10 壁电流分布、开槽是否扰乱场型及波导测量线背景。", "融入", ("content/knowledge/05-矩形波导工程计算/05-波导段反射驻波与匹配.md", "content/knowledge/05-矩形波导工程计算/99-自检清单与常见误区.md")),
        PageRange(31, 36, "波导激励与耦合：探针、环、小孔/缝，及定向耦合器背景。", "融入", ("content/knowledge/05-矩形波导工程计算/05-波导段反射驻波与匹配.md", "content/knowledge/07-实验测量与微波元件/04-定向耦合器与功率分配器.md")),
        PageRange(37, 37, "本讲概念总结：无界媒质与波导中的 λ、λg、vp、vg 关系。", "融入", ("content/knowledge/04-截止色散与速度/README.md", "content/knowledge/05-矩形波导工程计算/README.md")),
        PageRange(38, 38, "结束页。", "审计", ()),
    ),
    "circular-waveguide": (
        PageRange(1, 2, "封面、教学目的。", "审计", ()),
        PageRange(3, 3, "圆波导研究动机：对称、双极化、损耗与功率容量优势。", "融入", ("content/knowledge/06-圆波导同轴线微带线/01-圆波导模式与贝塞尔根.md", "content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md")),
        PageRange(4, 8, "柱坐标亥姆霍兹方程、分离变量、贝塞尔函数和横纵场关系。", "融入", ("content/knowledge/06-圆波导同轴线微带线/01-圆波导模式与贝塞尔根.md",)),
        PageRange(9, 16, "TM/TE 模边界条件、贝塞尔根、截止波长、TE11 主模和单模条件。", "融入", ("content/knowledge/06-圆波导同轴线微带线/01-圆波导模式与贝塞尔根.md",)),
        PageRange(17, 19, "TE11 场分布、极化简并、方圆过渡和 TE10 激励 TE11。", "融入", ("content/knowledge/06-圆波导同轴线微带线/01-圆波导模式与贝塞尔根.md", "content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md")),
        PageRange(20, 25, "TM01 与 TE01 的场结构、应用、壁电流和低损耗特性。", "融入", ("content/knowledge/06-圆波导同轴线微带线/01-圆波导模式与贝塞尔根.md", "content/knowledge/08-谐振器网络与课程综合/01-微波谐振器与谐振腔.md")),
        PageRange(26, 28, "EH 简并、极化简并和圆波导功率表达式。", "融入", ("content/knowledge/06-圆波导同轴线微带线/99-自检清单与常见误区.md",)),
        PageRange(29, 30, "贝塞尔函数背景和空白页。", "审计", ("content/knowledge/06-圆波导同轴线微带线/01-圆波导模式与贝塞尔根.md",)),
        PageRange(31, 31, "结束页。", "审计", ()),
    ),
    "coaxial-line": (
        PageRange(1, 2, "封面、教学目的。", "审计", ()),
        PageRange(3, 4, "硬同轴/软同轴结构、TEM 主模优点和高频限制。", "融入", ("content/knowledge/06-圆波导同轴线微带线/02-同轴线TEM与高阶模.md", "content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md")),
        PageRange(5, 12, "同轴 TEM 场分布、Er/Hφ、β、λ、vp、特性阻抗。", "融入", ("content/knowledge/06-圆波导同轴线微带线/02-同轴线TEM与高阶模.md",)),
        PageRange(13, 16, "同轴高阶 TE/TM 模、TM01、TE11 截止和简并关系。", "融入", ("content/knowledge/06-圆波导同轴线微带线/02-同轴线TEM与高阶模.md", "content/knowledge/06-圆波导同轴线微带线/99-自检清单与常见误区.md")),
        PageRange(17, 19, "同轴单模尺寸、功率容量、衰减常数和 b/a 折衷。", "融入", ("content/knowledge/06-圆波导同轴线微带线/02-同轴线TEM与高阶模.md", "content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md")),
        PageRange(20, 20, "结束页。", "审计", ()),
    ),
    "microwave-transmission-lines": (
        PageRange(1, 2, "封面、教学目的。", "审计", ()),
        PageRange(3, 5, "什么是微波传输线、为什么研究传输线、本征问题与激励问题。", "融入", ("content/knowledge/03-波导中的场与边界/00-从传输线到波导.md", "content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md")),
        PageRange(6, 10, "从 Maxwell 方程到亥姆霍兹方程、广义传输线理论。", "融入", ("content/knowledge/03-波导中的场与边界/02-纵向分量与分离变量.md",)),
        PageRange(11, 11, "TEM、TE、TM 模式定义和横向场正交性。", "融入", ("content/knowledge/03-波导中的场与边界/01-TEM-TE-TM波型.md",)),
        PageRange(12, 18, "规则金属波导假设、坐标/算子分离、传播因子和横向本征方程。", "融入", ("content/knowledge/03-波导中的场与边界/02-纵向分量与分离变量.md",)),
        PageRange(19, 20, "截止波数、截止波长、传播/截止条件。", "融入", ("content/knowledge/03-波导中的场与边界/03-金属边界与截止.md", "content/knowledge/04-截止色散与速度/00-从截止到色散.md")),
        PageRange(21, 23, "波动方程、模式和传输条件总结页。", "融入", ("content/knowledge/03-波导中的场与边界/99-自检清单与常见误区.md", "content/knowledge/04-截止色散与速度/99-自检清单与常见误区.md")),
        PageRange(24, 24, "结束页。", "审计", ()),
    ),
    "integrated-microwave-lines": (
        PageRange(1, 2, "封面、教学目标。", "审计", ()),
        PageRange(3, 6, "MIC、MMIC、MHMIC 背景和工程动机。", "提炼少量背景", ("content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md",)),
        PageRange(7, 8, "集成传输线特点及 TEM、准 TEM、非 TEM 分类。", "融入", ("content/knowledge/06-圆波导同轴线微带线/README.md", "content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md")),
        PageRange(9, 17, "带状线由同轴演化、TEM 主模、Zc、衰减、相速和波导波长。", "融入", ("content/knowledge/06-圆波导同轴线微带线/03-微带线准TEM与有效介电常数.md", "content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md")),
        PageRange(18, 20, "带状线单模、辐射损耗和匹配设计约束。", "融入", ("content/knowledge/06-圆波导同轴线微带线/04-从矩形到圆与微带的对照.md", "content/knowledge/06-圆波导同轴线微带线/99-自检清单与常见误区.md")),
        PageRange(21, 23, "微带线概念、不对称结构、有效介电常数等效思想。", "融入", ("content/knowledge/06-圆波导同轴线微带线/03-微带线准TEM与有效介电常数.md",)),
        PageRange(24, 27, "微带准 TEM 不是纯 TEM 的边界条件证明和纵向分量来源。", "融入", ("content/knowledge/06-圆波导同轴线微带线/03-微带线准TEM与有效介电常数.md", "content/knowledge/03-波导中的场与边界/01-TEM-TE-TM波型.md")),
        PageRange(28, 34, "微带特性阻抗、相速、有效介电常数、导体厚度修正、λg 和损耗。", "融入", ("content/knowledge/06-圆波导同轴线微带线/03-微带线准TEM与有效介电常数.md",)),
        PageRange(35, 39, "微带色散、高阶模、表面波模和强耦合工作频率限制。", "融入", ("content/knowledge/06-圆波导同轴线微带线/03-微带线准TEM与有效介电常数.md", "content/knowledge/06-圆波导同轴线微带线/99-自检清单与常见误区.md")),
        PageRange(40, 40, "耦合微带线及其在定向耦合器、滤波器和平衡变换中的用途。", "融入", ("content/knowledge/07-实验测量与微波元件/04-定向耦合器与功率分配器.md", "content/knowledge/08-谐振器网络与课程综合/03-常用微波元件网络化描述.md")),
        PageRange(41, 41, "结束页。", "审计", ()),
    ),
    "microwave-resonator": (
        PageRange(1, 2, "封面、教学目的。", "审计", ()),
        PageRange(3, 8, "LC 回路到微波谐振器、传输线型/非传输线型、分布储能和多模特征。", "融入", ("content/knowledge/08-谐振器网络与课程综合/01-微波谐振器与谐振腔.md",)),
        PageRange(9, 13, "谐振频率求法：相位法、电纳法、集中参数法、场叠加法。", "融入", ("content/knowledge/08-谐振器网络与课程综合/01-微波谐振器与谐振腔.md",)),
        PageRange(14, 20, "Q0、QL、外部 Q、耦合系数、等效电导和损耗表达。", "融入", ("content/knowledge/07-实验测量与微波元件/03-谐振器Q值与功率传输法.md", "content/knowledge/08-谐振器网络与课程综合/01-微波谐振器与谐振腔.md")),
        PageRange(21, 36, "矩形谐振腔 TE/TM 模、谐振频率、TE101 主模和 Q0。", "融入", ("content/knowledge/08-谐振器网络与课程综合/01-微波谐振器与谐振腔.md", "content/knowledge/05-矩形波导工程计算/01-模谱主模与简并.md")),
        PageRange(37, 52, "圆柱谐振腔 TE/TM 模、模式图、干扰模与耦合抑制。", "融入", ("content/knowledge/08-谐振器网络与课程综合/01-微波谐振器与谐振腔.md", "content/knowledge/06-圆波导同轴线微带线/01-圆波导模式与贝塞尔根.md")),
        PageRange(53, 58, "TE011、TE111、TM010 三个常用圆柱腔模式的优缺点。", "融入", ("content/knowledge/08-谐振器网络与课程综合/01-微波谐振器与谐振腔.md",)),
        PageRange(59, 67, "同轴谐振腔主模 TEM、λ/2 与 λ/4 结构、单模尺寸和 Q0。", "融入", ("content/knowledge/08-谐振器网络与课程综合/01-微波谐振器与谐振腔.md", "content/knowledge/06-圆波导同轴线微带线/02-同轴线TEM与高阶模.md")),
        PageRange(68, 69, "电容加载型同轴谐振腔及参考面电纳求谐振条件。", "融入", ("content/knowledge/08-谐振器网络与课程综合/01-微波谐振器与谐振腔.md",)),
        PageRange(70, 70, "结束页。", "审计", ()),
    ),
    "microwave-network": (
        PageRange(1, 2, "封面、教学目的。", "审计", ()),
        PageRange(3, 5, "微波网络概念、单/双/多端口分类、参考面和模式等效注意事项。", "融入", ("content/knowledge/08-谐振器网络与课程综合/02-微波网络基础与S参数.md", "content/knowledge/07-实验测量与微波元件/01-S参数与矢量网络分析仪.md")),
        PageRange(6, 15, "模式等效电压/电流、波导与双线等效、TE10 等效归一化。", "融入", ("content/knowledge/08-谐振器网络与课程综合/02-微波网络基础与S参数.md",)),
        PageRange(16, 22, "电路参量归一化、阻抗矩阵 Z、导纳矩阵 Y 及双端口例题。", "融入", ("content/knowledge/08-谐振器网络与课程综合/02-微波网络基础与S参数.md",)),
        PageRange(23, 31, "为什么改用波参量、a/b 归一化波、S 矩阵定义和互易/对称/无耗性质。", "融入", ("content/knowledge/07-实验测量与微波元件/01-S参数与矢量网络分析仪.md", "content/knowledge/08-谐振器网络与课程综合/02-微波网络基础与S参数.md")),
        PageRange(32, 34, "互易双端口 S 参数测量、短路/开路/匹配三点法和多点法。", "融入", ("content/knowledge/07-实验测量与微波元件/01-S参数与矢量网络分析仪.md", "content/knowledge/08-谐振器网络与课程综合/04-微波测量与课程综合.md")),
        PageRange(35, 35, "传输矩阵 T 与级联网络表达。", "融入", ("content/knowledge/08-谐振器网络与课程综合/02-微波网络基础与S参数.md", "content/knowledge/08-谐振器网络与课程综合/03-常用微波元件网络化描述.md")),
        PageRange(36, 38, "无可抽取文字的图页，仅作为视觉复核。", "审计", ()),
        PageRange(39, 39, "结束页。", "审计", ()),
    ),
}


KEYWORD_TAGS: tuple[str, ...] = (
    "长线",
    "分布参数",
    "特性阻抗",
    "反射系数",
    "驻波",
    "输入阻抗",
    "Smith",
    "导纳",
    "匹配",
    "支节",
    "矩形波导",
    "圆波导",
    "同轴线",
    "微带",
    "带状线",
    "准TEM",
    "TEM",
    "TE",
    "TM",
    "截止",
    "导波波长",
    "相速",
    "群速",
    "简并",
    "贝塞尔",
    "谐振",
    "品质因数",
    "Q",
    "网络",
    "S参数",
    "散射",
    "二端口",
    "耦合器",
    "功分器",
)


EXP2_IMAGES: dict[str, str] = {
    "fig_q_halfpower.png": "exp2-q-halfpower-method.webp",
    "fig_coupler_ports.png": "exp2-coupler-port-map.webp",
    "fig_wilkinson.png": "exp2-wilkinson-topology.webp",
    "fig_setup_resonator.png": "exp2-setup-resonator.webp",
    "fig_setup_coupler.png": "exp2-setup-coupler.webp",
    "fig_setup_divider.png": "exp2-setup-divider.webp",
    "image4.jpeg": "exp2-resonator-s21-curve.webp",
    "image5.jpeg": "exp2-resonator-bandwidth-search.webp",
    "image6.jpeg": "exp2-coupler-coupling-s21.webp",
    "image11.jpeg": "exp2-coupler-through-s21.webp",
    "image15.jpeg": "exp2-divider-port2-loss.webp",
    "image17.jpeg": "exp2-divider-port3-loss.webp",
    "image19.jpeg": "exp2-divider-isolation.webp",
}


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def page_title(text: str) -> str:
    lines = [normalize_text(line) for line in text.splitlines() if normalize_text(line)]
    for line in lines[:14]:
        if 2 <= len(line) <= 54 and not re.fullmatch(r"[\d\W_]+", line):
            return line[:54]
    return lines[0][:54] if lines else "无可抽取文字"


def page_tags(text: str) -> list[str]:
    lowered = text.lower()
    tags = [keyword for keyword in KEYWORD_TAGS if keyword.lower() in lowered]
    return tags[:10]


def page_score(text: str, keywords: Iterable[str]) -> int:
    normalized = normalize_text(text).lower()
    return sum(normalized.count(keyword.lower()) for keyword in keywords)


def validate_page_ranges(source: CourseSource, page_count: int) -> None:
    covered: set[int] = set()
    ranges = COURSE_PAGE_RANGES[source.slug]
    for item in ranges:
        covered.update(range(item.start, item.end + 1))
    expected = set(range(1, page_count + 1))
    missing = expected - covered
    extra = covered - expected
    if missing or extra:
        raise ValueError(f"{source.filename} page range mismatch: missing={sorted(missing)}, extra={sorted(extra)}")


def page_range_for(source: CourseSource, page_number: int) -> PageRange:
    for item in COURSE_PAGE_RANGES[source.slug]:
        if item.start <= page_number <= item.end:
            return item
    raise ValueError(f"No page range for {source.filename} page {page_number}")


def build_page_index(source: CourseSource, doc: fitz.Document) -> list[dict[str, object]]:
    validate_page_ranges(source, doc.page_count)
    pages: list[dict[str, object]] = []
    for index, page in enumerate(doc):
        raw = page.get_text("text") or ""
        clean = normalize_text(raw)
        plan = page_range_for(source, index + 1)
        pages.append(
            {
                "page": index + 1,
                "title": page_title(raw),
                "tags": page_tags(clean),
                "short_excerpt": clean[:160],
                "range_summary": plan.summary,
                "treatment": plan.treatment,
                "targets": list(plan.targets),
            }
        )
    return pages


def render_page_preview(doc: fitz.Document, page_index: int, output: Path) -> None:
    page = doc[page_index]
    zoom = 1.25
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    output.parent.mkdir(parents=True, exist_ok=True)
    image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    image = ImageOps.exif_transpose(image).convert("RGB")
    image.thumbnail((1280, 1280))
    image.save(output, "WEBP", quality=82, method=6)


def summarize_pdf(source: CourseSource) -> dict[str, object]:
    path = PDF_SOURCE_DIR / source.filename
    if not path.exists():
        return {
            "filename": source.filename,
            "missing": True,
            "slug": source.slug,
            "topic": source.topic,
            "stages": source.stages,
            "targets": list(source.targets),
        }

    doc = fitz.open(path)
    page_index = build_page_index(source, doc)
    page_infos: list[tuple[int, int, int, str]] = []
    total_chars = 0
    text_pages = 0
    for index, page in enumerate(doc):
        text = page.get_text("text") or ""
        clean = normalize_text(text)
        if clean:
            text_pages += 1
            total_chars += len(clean)
        page_infos.append((page_score(clean, source.keywords), len(clean), index, clean[:180]))

    best = max(page_infos, key=lambda item: (item[0], item[1]))
    preview_name = f"course-{source.slug}-source-preview.webp"
    preview_path = COURSE_IMAGE_DIR / preview_name
    render_page_preview(doc, best[2], preview_path)

    return {
        "filename": source.filename,
        "slug": source.slug,
        "topic": source.topic,
        "stages": source.stages,
        "targets": list(source.targets),
        "pages": doc.page_count,
        "text_pages": text_pages,
        "text_chars": total_chars,
        "selected_page": best[2] + 1,
        "selected_page_score": best[0],
        "selected_excerpt": best[3],
        "preview": f"assets/images/course/{preview_name}",
        "page_ranges": [
            {
                "pages": f"{item.start}-{item.end}" if item.start != item.end else str(item.start),
                "summary": item.summary,
                "treatment": item.treatment,
                "targets": list(item.targets),
            }
            for item in COURSE_PAGE_RANGES[source.slug]
        ],
        "page_index": page_index,
    }


def convert_image(src: Path, dest: Path, max_size: tuple[int, int] = (1800, 1400)) -> None:
    dest.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(src) as image:
        image = ImageOps.exif_transpose(image)
        if image.mode in {"RGBA", "LA"}:
            background = Image.new("RGB", image.size, "white")
            background.paste(image, mask=image.getchannel("A"))
            image = background
        else:
            image = image.convert("RGB")
        image.thumbnail(max_size)
        image.save(dest, "WEBP", quality=86, method=6)


def copy_exp2_images() -> list[dict[str, str]]:
    copied: list[dict[str, str]] = []
    media_dir = EXP2_SOURCE_DIR / "media"
    for source_name, dest_name in EXP2_IMAGES.items():
        src = media_dir / source_name
        if not src.exists():
            copied.append({"source": source_name, "missing": "true"})
            continue
        dest = EXP2_IMAGE_DIR / dest_name
        convert_image(src, dest)
        copied.append(
            {
                "source": f"sources/微波实验2/media/{source_name}",
                "published": f"assets/images/exp2/{dest_name}",
            }
        )
    return copied


def summarize_exp2_report() -> dict[str, object]:
    md_path = EXP2_SOURCE_DIR / "微波技术基础实验报告二（重构）.md"
    docx_path = EXP2_SOURCE_DIR / "微波技术基础实验报告二（参考）.docx"
    summary: dict[str, object] = {
        "markdown": str(md_path.relative_to(ROOT)) if md_path.exists() else "",
        "docx": str(docx_path.relative_to(ROOT)) if docx_path.exists() else "",
    }

    if md_path.exists():
        text = md_path.read_text(encoding="utf-8")
        summary["markdown_lines"] = len(text.splitlines())
        summary["markdown_headings"] = re.findall(r"^#{1,3}\s+(.+)$", text, flags=re.MULTILINE)

    if docx_path.exists():
        doc = Document(docx_path)
        summary["docx_paragraphs"] = sum(1 for paragraph in doc.paragraphs if paragraph.text.strip())
        summary["docx_tables"] = len(doc.tables)

    return summary


def write_audit(index: dict[str, object]) -> None:
    lines = [
        "# 来源抽取与融入审计",
        "",
        "更新时间：2026-06-03",
        "",
        "本文件记录 `sources/` 本地资料如何被消化进站点。原始 PDF/DOCX 不发布、不纳入版本控制；站点只使用重构后的 Markdown、审计摘要和精选 WebP 图片。",
        "",
        "页级复核详见 `docs/audit/COURSE_INTEGRATION_REVIEW.md`；完整逐页 JSON 索引见 `docs/audit/source_extraction_index.json`。JSON 每页只保存标题候选、关键词标签和短摘要，不保存完整课件正文。",
        "",
        "## 课件 PDF 映射",
        "",
        "| 源文件 | 页数 | 主题 | 融入阶段 | 目标页面 | 精选页 | 发布预览 | 采用内容 | 舍弃内容 | 复核状态 |",
        "|---|---:|---|---|---|---:|---|---|---|---|",
    ]
    for item in index["course_sources"]:
        targets = "<br>".join(item.get("targets", []))
        pages = item.get("pages", "缺失")
        selected = item.get("selected_page", "")
        preview = item.get("preview", "")
        adopted = "概念主线、核心公式、例题套路、可解释图像"
        discarded = "封面目录、重复板书、整页原文截图"
        review = "已抽取页码与候选图，人工重构入现有页"
        lines.append(
            f"| {item['filename']} | {pages} | {item['topic']} | {item['stages']} | {targets} | {selected} | `{preview}` | {adopted} | {discarded} | {review} |"
        )

    lines.extend(
        [
            "",
            "## 实验二资料映射",
            "",
            "实验二以 `sources/微波实验2/微波技术基础实验报告二（重构）.md` 为主输入，DOCX 只用于确认段落和表格结构。站点页面保留数据链，不发布原始报告文件。",
            "",
            "| 发布图片 | 来源 | 用途 |",
            "|---|---|---|",
        ]
    )
    for item in index["exp2_images"]:
        if item.get("missing"):
            lines.append(f"| 缺失 | {item['source']} | 待补 |")
            continue
        published = item["published"]
        source = item["source"]
        purpose = published.replace("assets/images/exp2/", "").replace(".webp", "")
        lines.append(f"| `{published}` | `{source}` | {purpose} |")

    report = index["exp2_report"]
    headings = report.get("markdown_headings", [])
    lines.extend(
        [
            "",
            "## 实验二报告结构",
            "",
            f"- Markdown 行数：{report.get('markdown_lines', 0)}",
            f"- DOCX 段落数：{report.get('docx_paragraphs', 0)}",
            f"- DOCX 表格数：{report.get('docx_tables', 0)}",
            f"- 主要标题：{'；'.join(headings[:16])}",
            "",
            "## 融入原则",
            "",
            "1. 只把课件中的概念、图像和例题结构转写成站点语言，不整页搬运 PPT 原文。",
            "2. 同一概念优先融入现有知识点页，不新增独立课件库。",
            "3. 实验数据必须写成“读数 → 换算 → 结论 → 误差来源”。",
            "4. 课件和实验二补强后必须更新交叉链接，保证能从知识点、实验页和作业页互相到达。",
        ]
    )
    AUDIT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_integration_review(index: dict[str, object]) -> None:
    lines = [
        "# 课件逐页理解与知识体系融入复核",
        "",
        "更新时间：2026-06-03",
        "",
        "本文件用于回答“每一张 PPT 应该如何理解、应该融入哪里”。处理原则是先逐页抽取主题信号，再把每页归入现有 `knowledge / experiments / solutions` 主线；封面、结束页、重复引入页和空白图页只保留审计记录，不发布为站点正文。",
        "",
        "完整逐页索引保存在 `docs/audit/source_extraction_index.json` 的 `course_sources[].page_index[]` 中。这里按连续页段列出人工复核后的融入决策，每个页段都覆盖原 PDF 页码，不遗漏页面。",
    ]

    for item in index["course_sources"]:
        lines.extend(
            [
                "",
                f"## {item['filename']}",
                "",
                f"- 页数：{item.get('pages', '缺失')}",
                f"- 主题判断：{item['topic']}",
                f"- 总体融入阶段：{item['stages']}",
                "",
                "| 页码 | 页级理解 | 处理 | 融入目标 |",
                "|---|---|---|---|",
            ]
        )
        for page_range in item.get("page_ranges", []):
            targets = "<br>".join(page_range.get("targets", [])) or "不进入正文"
            lines.append(
                f"| {page_range['pages']} | {page_range['summary']} | {page_range['treatment']} | {targets} |"
            )

    lines.extend(
        [
            "",
            "## 复核结论",
            "",
            "1. 10 份课件均已逐页覆盖；封面、教学目的、结束页、空白图页和重复铺垫页保留在审计中，不作为站点正文。",
            "2. 可教学化的内容全部落到现有知识点主线：传输线与 Smith 圆图落入 01-02；规则波导落入 03-05；圆波导、同轴线、微带/带状线落入 06；谐振器、网络与测量落入 07-08。",
            "3. 站点正文只吸收概念、公式前提、例题链、工程判断和必要图像；不发布原始 PDF，不新增课件库。",
        ]
    )
    INTEGRATION_REVIEW_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    COURSE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    EXP2_IMAGE_DIR.mkdir(parents=True, exist_ok=True)

    course_items = [summarize_pdf(source) for source in COURSE_SOURCES]
    exp2_images = copy_exp2_images()
    exp2_report = summarize_exp2_report()
    index = {
        "course_sources": course_items,
        "exp2_images": exp2_images,
        "exp2_report": exp2_report,
    }

    AUDIT_JSON.write_text(json.dumps(index, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    write_audit(index)
    write_integration_review(index)
    print(f"Wrote {AUDIT_MD.relative_to(ROOT)}")
    print(f"Wrote {AUDIT_JSON.relative_to(ROOT)}")
    print(f"Wrote {INTEGRATION_REVIEW_MD.relative_to(ROOT)}")
    print(f"Course previews: {len(course_items)}")
    print(f"Experiment 2 images: {sum(1 for item in exp2_images if not item.get('missing'))}")


if __name__ == "__main__":
    main()
